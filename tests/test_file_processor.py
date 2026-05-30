import pytest

from app.utils.file_processor import FileProcessor
from unittest.mock import patch, MagicMock


class TestFileProcessor:
    """文件处理器测试"""

    def setup_method(self):
        self.processor = FileProcessor()

    def test_read_txt(self):
        """测试读取txt文件"""
        content = "Hello, World!\nThis is a test."
        file_bytes = content.encode("utf-8")

        result = FileProcessor.read_txt(file_bytes)
        assert result == content

    def test_read_md_preserves_structure(self):
        """测试 MD 直接保留原文，不丢失排版结构"""
        md_content = """# 第一章 概述

这是概述内容。

## 1.1 背景

背景说明。

### 1.1.1 详细说明

详细内容。

| 列1 | 列2 |
|-----|-----|
| A   | B   |

## 1.2 目标

目标说明。"""

        file_bytes = md_content.encode("utf-8")
        result = FileProcessor.read_md(file_bytes)

        # 应该保留所有标题标记
        assert "# 第一章 概述" in result
        assert "## 1.1 背景" in result
        assert "### 1.1.1 详细说明" in result
        # 应该保留表格
        assert "| 列1 | 列2 |" in result
        # 不应该被转成纯文本丢失结构
        assert "## 1.2 目标" in result

    def test_read_docx_preserves_headings(self):
        """测试 DOCX 结构化提取保留标题层级"""
        # 构造一个简单的 docx 文件（用 python-docx 生成）
        from docx import Document
        from docx.shared import Pt
        import io

        doc = Document()
        doc.add_heading("第一章 概述", level=1)
        doc.add_paragraph("这是概述内容。")
        doc.add_heading("1.1 背景", level=2)
        doc.add_paragraph("背景说明。")

        buf = io.BytesIO()
        doc.save(buf)
        file_bytes = buf.getvalue()

        result = FileProcessor.read_docx(file_bytes)

        # 应该转为 Markdown 标题格式
        assert "# 第一章 概述" in result
        assert "## 1.1 背景" in result
        assert "概述内容" in result
        assert "背景说明" in result

    def test_read_docx_preserves_tables(self):
        """测试 DOCX 结构化提取保留表格"""
        from docx import Document
        import io

        doc = Document()
        doc.add_heading("数据表", level=1)
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "年龄"
        table.cell(1, 0).text = "张三"
        table.cell(1, 1).text = "25"

        buf = io.BytesIO()
        doc.save(buf)
        file_bytes = buf.getvalue()

        result = FileProcessor.read_docx(file_bytes)

        assert "# 数据表" in result
        # 表格应该转为 Markdown 格式
        assert "| 姓名 | 年龄 |" in result
        assert "| --- | --- |" in result
        assert "| 张三 | 25 |" in result

    def test_split_md_by_headers(self):
        """测试 Markdown 按标题层级切片"""
        md_text = """# 第一章

第一章的内容比较长，包含了多个段落的详细信息。

## 1.1 小节

这是 1.1 小节的内容，讲述了具体的细节。

## 1.2 另一个小节

这是 1.2 小节的内容，提供了更多说明。"""

        chunks = self.processor.split_text(md_text, chunk_size=500, overlap=50)

        # 应该按标题层级切片，至少有 3 个 chunk
        assert len(chunks) >= 3
        # 每个 chunk 应该包含对应的标题内容
        chunks_text = "\n".join(chunks)
        assert "第一章" in chunks_text
        assert "1.1" in chunks_text
        assert "1.2" in chunks_text

    def test_split_md_no_headers_fallback(self):
        """测试无标题的 Markdown 退回字符分割"""
        md_text = "这是一段没有标题的纯文本内容。" * 50

        chunks = self.processor.split_text(md_text, chunk_size=500, overlap=50)

        assert len(chunks) >= 2
        assert all(len(c) <= 500 for c in chunks)

    def test_split_text_basic(self):
        """测试基础文本分块（纯文本，无 Markdown 标题）"""
        text = "a" * 1000
        chunks = self.processor.split_text(text, chunk_size=500, overlap=50)

        assert len(chunks) >= 2
        assert all(len(c) <= 500 for c in chunks)

    def test_split_text_with_overlap(self):
        """测试带重叠的文本分块"""
        text = "abcdefghij" * 10  # 100 chars
        chunks = self.processor.split_text(text, chunk_size=50, overlap=10)

        assert len(chunks) >= 2
        assert all(len(c) <= 50 for c in chunks)

    def test_split_text_small_text(self):
        """测试小块文本处理"""
        text = "short text"
        chunks = self.processor.split_text(text, chunk_size=500, overlap=50)

        assert len(chunks) == 1
        assert text in chunks[0]

    def test_process_txt_file(self):
        """测试处理txt文件"""
        filename = "test.txt"
        content = "This is test content."
        file_bytes = content.encode("utf-8")

        result = FileProcessor.process(filename, file_bytes)
        assert result == content

    def test_process_md_file(self):
        """测试处理 md 文件"""
        filename = "test.md"
        content = "# Title\n\nSome content"
        file_bytes = content.encode("utf-8")

        result = FileProcessor.process(filename, file_bytes)
        assert result == content

    def test_process_unsupported_file_type(self):
        """测试不支持的文件类型"""
        filename = "test.xyz"
        file_bytes = b"some content"

        with pytest.raises(ValueError) as exc_info:
            FileProcessor.process(filename, file_bytes)

        assert "Unsupported file type" in str(exc_info.value)

    def test_process_file_case_insensitive(self):
        """测试文件扩展名大小写不敏感"""
        filename = "test.TXT"
        content = "Test content"
        file_bytes = content.encode("utf-8")

        result = FileProcessor.process(filename, file_bytes)
        assert result == content

    def test_split_docx_with_headings(self):
        """测试 DOCX 文件按标题切片的完整流程"""
        from docx import Document
        import io

        doc = Document()
        doc.add_heading("项目介绍", level=1)
        doc.add_paragraph("这是一个关于智能路由的项目。")
        doc.add_heading("技术架构", level=2)
        doc.add_paragraph("项目采用 FastAPI 框架。")
        doc.add_heading("功能特性", level=2)
        doc.add_paragraph("支持多种文件格式的解析。")

        buf = io.BytesIO()
        doc.save(buf)
        file_bytes = buf.getvalue()

        # 先解析
        text = FileProcessor.read_docx(file_bytes)
        # 再切片
        chunks = self.processor.split_text(text, chunk_size=500, overlap=50)

        # 应该按标题切片
        assert len(chunks) >= 3
        chunks_text = "\n".join(chunks)
        assert "项目介绍" in chunks_text
        assert "技术架构" in chunks_text
        assert "功能特性" in chunks_text

    def test_split_long_section_sub_split(self):
        """测试超长章节会被二次字符分割"""
        # 构造一个有标题但内容超长的 Markdown
        md_text = "# 长章节\n\n" + "这是一段很长的内容。" * 200  # ~2200 字符

        chunks = self.processor.split_text(md_text, chunk_size=500, overlap=50)

        # 应该被二次分割为多个 chunk
        assert len(chunks) >= 4
        # 每个 chunk 都不应超长
        for chunk in chunks:
            assert len(chunk) <= 500

    def test_read_pdf_success(self):
        """测试 PDF 通过飞桨 OCR API 成功解析"""
        fake_job_id = "job_test_123"
        mock_submit_resp = MagicMock()
        mock_submit_resp.json.return_value = {"id": fake_job_id}
        mock_submit_resp.raise_for_status = MagicMock()

        mock_poll_resp = MagicMock()
        mock_poll_resp.json.return_value = {
            "status": "completed",
            "result": {
                "markdown": "# 测试文档\n\n这是 PDF 解析出的内容。"
            }
        }
        mock_poll_resp.raise_for_status = MagicMock()

        mock_client = MagicMock()
        mock_client.post.return_value = mock_submit_resp
        mock_client.get.return_value = mock_poll_resp
        mock_client.__enter__ = MagicMock(return_value=mock_client)
        mock_client.__exit__ = MagicMock(return_value=False)

        with patch("app.utils.file_processor.httpx.Client", return_value=mock_client):
            result = FileProcessor.read_pdf(b"fake pdf bytes")

        assert "# 测试文档" in result
        assert "PDF 解析出的内容" in result
        mock_client.post.assert_called_once()

    def test_read_pdf_no_token(self):
        """测试未配置 PaddleOCR token 时报错"""
        with patch("app.utils.file_processor.config") as mock_config:
            mock_config.paddle_ocr.token = ""
            with pytest.raises(ValueError, match="PaddleOCR token 未配置"):
                FileProcessor.read_pdf(b"fake pdf bytes")

    def test_read_pdf_job_failed(self):
        """测试 PaddleOCR 任务失败时报错"""
        mock_submit_resp = MagicMock()
        mock_submit_resp.json.return_value = {"id": "job_fail"}
        mock_submit_resp.raise_for_status = MagicMock()

        mock_poll_resp = MagicMock()
        mock_poll_resp.json.return_value = {
            "status": "failed",
            "error": "文件格式不支持"
        }
        mock_poll_resp.raise_for_status = MagicMock()

        mock_client = MagicMock()
        mock_client.post.return_value = mock_submit_resp
        mock_client.get.return_value = mock_poll_resp
        mock_client.__enter__ = MagicMock(return_value=mock_client)
        mock_client.__exit__ = MagicMock(return_value=False)

        with patch("app.utils.file_processor.httpx.Client", return_value=mock_client):
            with pytest.raises(RuntimeError, match="PaddleOCR 任务失败"):
                FileProcessor.read_pdf(b"fake pdf bytes")
