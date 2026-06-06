import io
import json
import time
from pathlib import Path

import httpx
from docx import Document
from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter

from app.config import config


class FileProcessor:
    """文件解析器，支持 txt, pdf, md, docx。

    解析策略：所有格式统一转为结构化 Markdown，再按标题层级语义切片。
    - md: 直接保留原文（本身已是 Markdown）
    - docx: 结构化提取标题层级 + 表格 → Markdown
    - pdf: 飞桨 PaddleOCR-VL 云端 API 识别 → Markdown
    - txt: 纯文本无结构，直接读取
    """

    def __init__(self):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.rag.chunk_size,
            chunk_overlap=config.rag.chunk_overlap,
        )

    @staticmethod
    def read_txt(file: bytes) -> str:
        return file.decode("utf-8")

    @staticmethod
    def read_pdf(file: bytes) -> str:
        """PDF → 飞桨 PaddleOCR-VL 云端 API → 结构化 Markdown"""
        paddle_cfg = config.paddle_ocr

        if not paddle_cfg.token:
            raise ValueError("PaddleOCR token 未配置，请在 config.yaml 或环境变量 PADDLEOCR_ACCESS_TOKEN 中设置")

        headers = {"Authorization": f"Bearer {paddle_cfg.token}"}
        files = {"file": ("document.pdf", file, "application/pdf")}
        data = {"model": paddle_cfg.model}

        with httpx.Client(timeout=60.0) as client:
            # 1. 提交 OCR 任务
            resp = client.post(paddle_cfg.endpoint, headers=headers, files=files, data=data)
            resp.raise_for_status()
            job_data = resp.json()
            job_data_inner = job_data.get("data", job_data)
            job_id = job_data_inner.get("jobId") or job_data_inner.get("id") or job_data_inner.get("job_id")
            if not job_id:
                raise RuntimeError(f"PaddleOCR 提交任务失败，响应: {job_data}")

            # 2. 轮询任务状态
            status_url = f"{paddle_cfg.endpoint}/{job_id}"
            max_retries = 120
            poll_interval = 2

            for _ in range(max_retries):
                resp = client.get(status_url, headers=headers)
                resp.raise_for_status()
                status_data = resp.json()
                # API 响应格式: {"code":0, "msg":"...", "data":{"state":"done", "resultUrl":{"jsonUrl":"..."}}}
                data = status_data.get("data", status_data)
                # API 用 "state" 而非 "status"，兼容两种命名
                state = data.get("state") or data.get("status") or ""

                if state in ("completed", "success", "done"):
                    # 结果不在响应内，需要从 resultUrl.jsonUrl 下载
                    result_url_info = data.get("resultUrl", {})
                    json_url = result_url_info.get("jsonUrl") if isinstance(result_url_info, dict) else None

                    if json_url:
                        result_resp = client.get(json_url, timeout=30.0)
                        result_resp.raise_for_status()
                        result_text = result_resp.text
                        # jsonUrl 可能返回 JSONL（每行一个 JSON 对象），而非标准 JSON
                        try:
                            result = result_resp.json()
                        except Exception:
                            # 尝试逐行解析 JSONL
                            result = []
                            for line in result_text.strip().splitlines():
                                line = line.strip()
                                if line:
                                    try:
                                        result.append(json.loads(line))
                                    except Exception:
                                        pass
                        # 解析 PaddleOCR-VL 结果结构
                        markdown_text = FileProcessor._extract_markdown_from_paddle_result(result)
                    else:
                        # 降级：尝试从 data 内部直接取结果
                        result = data.get("result", data)
                        markdown_text = (
                            result.get("markdown")
                            or result.get("content")
                            or result.get("text")
                            or ""
                        )
                        if not markdown_text and isinstance(result, str):
                            markdown_text = result

                    if not markdown_text:
                        raise RuntimeError(f"PaddleOCR 任务完成但结果为空，响应: {status_data}")
                    return markdown_text

                if state in ("failed", "error"):
                    error_msg = data.get("error") or data.get("message") or status_data.get("error") or status_data.get("message") or "未知错误"
                    raise RuntimeError(f"PaddleOCR 任务失败: {error_msg}")

                time.sleep(poll_interval)

            raise TimeoutError("PaddleOCR 任务超时（超过 4 分钟未完成）")

    @staticmethod
    def read_md(file: bytes) -> str:
        """直接保留 Markdown 原文，不做 HTML 转换，保留排版结构"""
        return file.decode("utf-8")

    @staticmethod
    def read_docx(file: bytes) -> str:
        """结构化提取 DOCX，保留标题层级和表格，输出 Markdown 格式"""
        doc = Document(io.BytesIO(file))
        lines = []

        # 建立 paragraph 和 table 与 body element 的对应关系
        para_index = 0
        table_index = 0

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                if para_index >= len(doc.paragraphs):
                    continue
                para = doc.paragraphs[para_index]
                para_index += 1

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""

                # 根据标题样式添加 # 前缀
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.replace("Heading", "").strip())
                    except ValueError:
                        level = 1
                    lines.append(f"{'#' * level} {text}")
                elif style_name.startswith("Title"):
                    lines.append(f"# {text}")
                else:
                    lines.append(text)

            elif tag == "tbl":
                if table_index >= len(doc.tables):
                    continue
                table = doc.tables[table_index]
                table_index += 1

                md_table = FileProcessor._table_to_markdown(table)
                if md_table:
                    lines.append(md_table)

        return "\n\n".join(lines)

    @staticmethod
    def _table_to_markdown(table) -> str:
        """将 docx 表格转为 Markdown 格式"""
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return ""

        # 去重合并单元格（docx 会对合并单元格重复）
        # 简单处理：直接输出
        col_count = len(rows_data[0])
        lines = []
        for i, row in enumerate(rows_data):
            # 补齐列数
            padded = row + [""] * (col_count - len(row))
            lines.append("| " + " | ".join(padded[:col_count]) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * col_count) + " |")

        return "\n".join(lines)

    @staticmethod
    def _extract_markdown_from_paddle_result(result) -> str:
        """从 PaddleOCR-VL JSONL 结果中提取 Markdown 文本。

        结果结构：每行 JSON 为一页 PDF 的结果
        {
          "result": {
            "layoutParsingResults": [{
              "prunedResult": {
                "parsing_res_list": [
                  {"block_label": "title", "block_content": "...", ...},
                  {"block_label": "text", "block_content": "...", ...},
                  ...
                ]
              }
            }]
          }
        }
        """
        # 标签 → Markdown 格式映射
        label_heading_map = {
            "title": "#",
            "header": "##",
            "section_header": "###",
        }
        # 需要跳过的标签（页眉页脚等噪音）
        skip_labels = {"header_image", "footer_image", "footer", "number", "footnote"}

        pages = []

        # result 可能是 list（JSONL）或 dict
        items = result if isinstance(result, list) else [result]

        for item in items:
            if not isinstance(item, dict):
                continue
            inner = item.get("result", item)
            if not isinstance(inner, dict):
                continue
            layouts = inner.get("layoutParsingResults", [])
            if not isinstance(layouts, list):
                layouts = [layouts]

            for layout in layouts:
                if not isinstance(layout, dict):
                    continue
                pruned = layout.get("prunedResult", layout)
                if not isinstance(pruned, dict):
                    continue
                res_list = pruned.get("parsing_res_list", [])
                if not isinstance(res_list, list):
                    continue

                page_lines = []
                for block in res_list:
                    if not isinstance(block, dict):
                        continue
                    label = block.get("block_label", "")
                    content = block.get("block_content", "")
                    if not content or label in skip_labels:
                        continue

                    if label in label_heading_map:
                        prefix = label_heading_map[label]
                        page_lines.append(f"{prefix} {content}")
                    elif label == "table":
                        # 表格内容可能已经是 Markdown 格式
                        page_lines.append(content)
                    elif label in ("figure", "figure_caption", "image"):
                        page_lines.append(content)
                    else:
                        # 普通文本块
                        page_lines.append(content)

                if page_lines:
                    pages.append("\n\n".join(page_lines))

        return "\n\n---\n\n".join(pages)

    @classmethod
    def process(cls, filename: str, file_bytes: bytes) -> str:
        ext = Path(filename).suffix.lower()

        handlers = {
            ".txt": cls.read_txt,
            ".pdf": cls.read_pdf,
            ".md": cls.read_md,
            ".docx": cls.read_docx,
        }

        if ext not in handlers:
            raise ValueError(f"Unsupported file type: {ext}")

        return handlers[ext](file_bytes)

    def split_text(self, text: str, chunk_size: int | None = None, overlap: int | None = None) -> list[str]:
        """统一切片逻辑：先按 Markdown 标题层级语义切片，超长 chunk 再二次字符分割"""
        effective_size = chunk_size or config.rag.chunk_size
        effective_overlap = overlap or config.rag.chunk_overlap

        # 第一步：按 Markdown 标题层级切片（保留语义边界）
        headers_to_split_on = [
            ("#", "h1"),
            ("##", "h2"),
            ("###", "h3"),
        ]
        md_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)

        try:
            md_chunks = md_splitter.split_text(text)
        except Exception:
            # 如果 Markdown 解析失败（如纯文本），退回字符分割
            char_splitter = RecursiveCharacterTextSplitter(
                chunk_size=effective_size,
                chunk_overlap=effective_overlap,
            )
            return char_splitter.split_text(text)

        # 如果没有标题结构（md_chunks 只有一个且无标题 metadata），退回字符分割
        if len(md_chunks) == 1 and not md_chunks[0].metadata:
            char_splitter = RecursiveCharacterTextSplitter(
                chunk_size=effective_size,
                chunk_overlap=effective_overlap,
            )
            return char_splitter.split_text(text)

        # 第二步：对超长 chunk 做二次字符分割
        char_splitter = RecursiveCharacterTextSplitter(
            chunk_size=effective_size,
            chunk_overlap=effective_overlap,
        )

        final_chunks = []
        for chunk in md_chunks:
            header_prefix = self._build_header_prefix(chunk.metadata)

            # 判断是否二次分割时，只看 page_content 本身长度
            # 标题前缀是为了给 chunk 提供上下文，不应成为触发二次分割的因素
            if len(chunk.page_content) > effective_size:
                # 为标题前缀预留空间，避免子 chunk 加上标题后超限
                reserved = len(header_prefix) + 2 if header_prefix else 0
                sub_size = max(effective_size - reserved, 100)
                sub_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=sub_size,
                    chunk_overlap=effective_overlap,
                )
                sub_chunks = sub_splitter.split_text(chunk.page_content)
                for sub in sub_chunks:
                    content = f"{header_prefix}\n{sub}" if header_prefix else sub
                    final_chunks.append(content)
            else:
                content = f"{header_prefix}\n{chunk.page_content}" if header_prefix else chunk.page_content
                final_chunks.append(content)

        return final_chunks

    @staticmethod
    def _build_header_prefix(metadata: dict) -> str:
        """从 MarkdownHeaderTextSplitter 的 metadata 中重建标题行"""
        header_map = {"h1": "#", "h2": "##", "h3": "###"}
        parts = []
        for level in ("h1", "h2", "h3"):
            if level in metadata:
                parts.append(f"{header_map[level]} {metadata[level]}")
        return "\n".join(parts)
